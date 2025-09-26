#!/usr/bin/env python3
"""
간단한 Word 템플릿 파일을 생성하는 스크립트
"""

from docx import Document
from docx.shared import Inches

def create_simple_template():
    """간단한 견적서 템플릿 생성"""
    
    # 새 문서 생성
    doc = Document()
    
    # 제목
    doc.add_heading('LRQA ISO 인증 견적서', 0)
    
    # 견적서 정보
    doc.add_paragraph('일자 : {quotation_date}')
    doc.add_paragraph('김 기석 대리')
    doc.add_paragraph('No. {quotation_number}')
    
    # 회사 정보
    doc.add_heading('회사 정보', level=1)
    doc.add_paragraph('회사명: {client_name}')
    doc.add_paragraph('주소: {client_address}')
    doc.add_paragraph('담당자: {contact_person}')
    doc.add_paragraph('이메일: {contact_email}')
    doc.add_paragraph('전화: {contact_phone}')
    
    # 인증 표준
    doc.add_heading('인증 표준', level=1)
    doc.add_paragraph('표준: {standards_text}')
    doc.add_paragraph('사업장 수: {total_sites}')
    doc.add_paragraph('직원 수: {total_employees}')
    
    # 견적 정보
    doc.add_heading('견적 정보', level=1)
    doc.add_paragraph('총 심사 일수: {total_audit_days} days')
    doc.add_paragraph('총 비용 (여행비 포함): {total_cost_with_travel_formatted}원')
    
    # ISO 14001 정보
    doc.add_paragraph('ISO 14001 감시심사 일수: {iso14001_surveillance_days} days')
    doc.add_paragraph('ISO 14001 1단계+2단계 일수: {iso14001_stage1_2_days} days')
    doc.add_paragraph('ISO 14001 1단계+2단계 비용: {iso14001_stage1_2_cost_formatted}원')
    doc.add_paragraph('여행비: {travel_expense_formatted}원')
    
    # 조건부 표시
    doc.add_paragraph('ISO 9001: {has_iso9001}')
    doc.add_paragraph('ISO 14001: {has_iso14001}')
    doc.add_paragraph('ISO 45001: {has_iso45001}')
    
    # 파일 저장
    doc.save('public/templates/LRQA_quotation.docx')
    print('간단한 템플릿 파일 생성 완료: public/templates/LRQA_quotation.docx')

if __name__ == "__main__":
    create_simple_template()
