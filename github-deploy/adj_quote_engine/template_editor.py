"""
LRQA 견적서 템플릿 편집 도구

이 도구는 LRQA_quotation_template.docx 파일에 Jinja2 변수를 자동으로 추가하는 도구입니다.
"""

import os
import shutil
from docxtpl import DocxTemplate
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


class TemplateEditor:
    """템플릿 편집기"""
    
    def __init__(self):
        self.template_path = os.path.join(os.path.dirname(__file__), 'templates', 'LRQA_quotation_template.docx')
        self.backup_path = os.path.join(os.path.dirname(__file__), 'templates', 'LRQA_quotation_template_backup.docx')
    
    def create_sample_template(self):
        """샘플 템플릿 생성 (변수 포함)"""
        try:
            # 백업 생성
            if os.path.exists(self.template_path):
                shutil.copy2(self.template_path, self.backup_path)
                print(f"백업 파일 생성: {self.backup_path}")
            
            # 새 템플릿 생성
            doc = Document()
            
            # 제목
            title = doc.add_heading('ISO 인증심사 견적서', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # 견적서 정보
            doc.add_heading('견적서 정보', level=1)
            doc.add_paragraph(f"견적서 번호: {{ quotation_number }}")
            doc.add_paragraph(f"작성일: {{ quotation_date }}")
            doc.add_paragraph(f"유효기간: {{ valid_until }}")
            
            # 고객사 정보
            doc.add_heading('고객사 정보', level=1)
            doc.add_paragraph(f"회사명: {{ client_name }} ({{ client_name_en }})")
            doc.add_paragraph(f"주소: {{ client_address }}")
            doc.add_paragraph(f"담당자: {{ contact_person }}")
            doc.add_paragraph(f"이메일: {{ contact_email }}")
            doc.add_paragraph(f"전화번호: {{ contact_phone }}")
            
            # 신청 표준
            doc.add_heading('신청 표준', level=1)
            doc.add_paragraph(f"적용 표준: {{ standards_text }}")
            
            # 사업장 정보
            doc.add_heading('사업장 정보', level=1)
            doc.add_paragraph(f"총 사업장 수: {{ total_sites }}개")
            doc.add_paragraph(f"총 직원 수: {{ total_employees }}명")
            
            # 사업장 목록
            doc.add_paragraph("사업장 목록:")
            doc.add_paragraph("{% for site in sites %}")
            doc.add_paragraph("{{ site.number }}. {{ site.name }}")
            doc.add_paragraph("   주소: {{ site.address }}")
            doc.add_paragraph("   직원수: {{ site.headcount }}명")
            doc.add_paragraph("   적용표준: {{ site.standards }}")
            doc.add_paragraph("   주요활동: {{ site.activities }}")
            doc.add_paragraph("")
            doc.add_paragraph("{% endfor %}")
            
            # 직원 구성
            doc.add_heading('직원 구성', level=1)
            doc.add_paragraph(f"총 직원 수: {{ employee_breakdown.total }}명")
            doc.add_paragraph(f"정규직: {{ employee_breakdown.permanent }}명")
            doc.add_paragraph(f"비정규직: {{ employee_breakdown.temporary }}명")
            doc.add_paragraph(f"협력업체: {{ employee_breakdown.contractors }}명")
            
            # 견적 요약
            doc.add_heading('견적 요약', level=1)
            doc.add_paragraph(f"총 심사일수: {{ total_audit_days }} mandays")
            doc.add_paragraph(f"서브토탈: ₩{{ subtotal | int | format_currency }}")
            doc.add_paragraph(f"VAT (10%): ₩{{ vat_amount | int | format_currency }}")
            doc.add_paragraph(f"총 견적 금액: ₩{{ total_cost | int | format_currency }}")
            
            # 표준별 상세
            doc.add_heading('표준별 상세', level=1)
            doc.add_paragraph("{% for detail in quotation_details %}")
            doc.add_paragraph("{{ detail.standard_name }} ({{ detail.standard }})")
            doc.add_paragraph("   ENP: {{ detail.enp }}명")
            doc.add_paragraph("   복잡도: {{ detail.complexity }}")
            doc.add_paragraph("   Stage1: {{ detail.stage1_days }}일 (₩{{ detail.stage1_cost | int | format_currency }})")
            doc.add_paragraph("   Stage2: {{ detail.stage2_days }}일 (₩{{ detail.stage2_cost | int | format_currency }})")
            doc.add_paragraph("   Surveillance: {{ detail.surveillance_days }}일 (₩{{ detail.surveillance_cost | int | format_currency }})")
            doc.add_paragraph("   Recert: {{ detail.recert_days }}일 (₩{{ detail.recert_cost | int | format_currency }})")
            doc.add_paragraph("   소계: {{ detail.total_days }}일 (₩{{ detail.total_cost | int | format_currency }})")
            doc.add_paragraph("")
            doc.add_paragraph("{% endfor %}")
            
            # 할인 정보
            doc.add_heading('할인 정보', level=1)
            doc.add_paragraph(f"통합심사 여부: {{ is_integrated }}")
            doc.add_paragraph(f"통합심사 할인: {{ integration_discount }}%")
            doc.add_paragraph(f"원격심사 비율: {{ remote_audit_ratio }}%")
            doc.add_paragraph(f"원격심사 할인: {{ remote_discount }}%")
            
            # 가정 사항
            doc.add_heading('가정 사항', level=1)
            doc.add_paragraph("{% for assumption in assumptions %}")
            doc.add_paragraph("- {{ assumption }}")
            doc.add_paragraph("{% endfor %}")
            
            # 근거 사항
            doc.add_heading('근거 사항', level=1)
            doc.add_paragraph("{% for justification in justification %}")
            doc.add_paragraph("- {{ justification }}")
            doc.add_paragraph("{% endfor %}")
            
            # 작성자 정보
            doc.add_heading('작성자 정보', level=1)
            doc.add_paragraph(f"작성자: {{ prepared_by }}")
            doc.add_paragraph(f"소속: {{ prepared_title }}")
            doc.add_paragraph(f"생성일시: {{ created_at }}")
            
            # 문서 저장
            doc.save(self.template_path)
            print(f"샘플 템플릿 생성 완료: {self.template_path}")
            
            return True
            
        except Exception as e:
            print(f"샘플 템플릿 생성 실패: {str(e)}")
            return False
    
    def restore_backup(self):
        """백업 파일 복원"""
        try:
            if os.path.exists(self.backup_path):
                shutil.copy2(self.backup_path, self.template_path)
                print(f"백업 파일 복원 완료: {self.template_path}")
                return True
            else:
                print("백업 파일이 존재하지 않습니다.")
                return False
        except Exception as e:
            print(f"백업 파일 복원 실패: {str(e)}")
            return False


def main():
    """메인 함수"""
    editor = TemplateEditor()
    
    print("LRQA 견적서 템플릿 편집 도구")
    print("1. 샘플 템플릿 생성")
    print("2. 백업 파일 복원")
    print("3. 종료")
    
    choice = input("선택하세요 (1-3): ").strip()
    
    if choice == "1":
        editor.create_sample_template()
    elif choice == "2":
        editor.restore_backup()
    elif choice == "3":
        print("종료합니다.")
    else:
        print("잘못된 선택입니다.")


if __name__ == "__main__":
    main()
