from flask import Flask, request, send_file, jsonify
from flask_cors import CORS
import tempfile
import os
from datetime import datetime
from io import BytesIO
from docx import Document

app = Flask(__name__)
CORS(app)

def create_simple_quotation_docx(application_data):
    """간단한 견적서 Word 문서 생성"""
    
    # 새 문서 생성
    doc = Document()
    
    # 제목
    title = doc.add_heading('LRQA 견적서', 0)
    title.alignment = 1  # 가운데 정렬
    
    # 회사 정보
    doc.add_heading('회사 정보', level=1)
    doc.add_paragraph(f"회사명: {application_data.get('법인명(국문)', '알 수 없음')}")
    doc.add_paragraph(f"영문명: {application_data.get('법인명(영문)', 'Unknown')}")
    doc.add_paragraph(f"주소: {application_data.get('본사주소', '서울시 강남구')}")
    doc.add_paragraph(f"담당자: {application_data.get('담당자명', '알 수 없음')}")
    doc.add_paragraph(f"이메일: {application_data.get('담당자이메일', 'unknown@example.com')}")
    doc.add_paragraph(f"전화: {application_data.get('담당자전화', '010-0000-0000')}")
    
    # 신청 표준
    doc.add_heading('신청 표준', level=1)
    doc.add_paragraph(f"ISO 표준: {application_data.get('ISO표준', 'ISO 9001')}")
    
    # 조직 규모
    doc.add_heading('조직 규모', level=1)
    doc.add_paragraph(f"총 직원수: {application_data.get('총직원수', '30')}명")
    doc.add_paragraph(f"비정규직수: {application_data.get('비정규직수', '0')}명")
    doc.add_paragraph(f"협력업체 직원수: {application_data.get('협력업체직원수', '0')}명")
    doc.add_paragraph(f"교대근무자수: {application_data.get('교대근무자수', '0')}명")
    
    # 견적 정보 (간단한 계산)
    doc.add_heading('견적 정보', level=1)
    
    # 기본 심사일수 계산 (간단한 로직)
    total_employees = int(application_data.get('총직원수', 30))
    iso_standards = application_data.get('ISO표준', 'ISO 9001')
    standard_count = len([s for s in ['ISO 9001', 'ISO 14001', 'ISO 45001'] if s in iso_standards])
    
    # 기본 심사일수 계산 (간단한 공식)
    base_days = max(2, total_employees // 20) * standard_count
    total_days = base_days + (standard_count - 1) * 0.5  # 다중 표준 할인
    
    doc.add_paragraph(f"계산된 심사일수: {total_days:.1f}일")
    doc.add_paragraph(f"1일 단가: 1,400,000원")
    
    total_cost = total_days * 1400000
    vat = total_cost * 0.1
    final_cost = total_cost + vat
    
    doc.add_paragraph(f"심사비용: {total_cost:,.0f}원")
    doc.add_paragraph(f"부가세 (10%): {vat:,.0f}원")
    doc.add_paragraph(f"총 견적금액: {final_cost:,.0f}원")
    
    # 견적서 번호
    doc.add_heading('견적서 정보', level=1)
    quotation_number = f"LRQA-{datetime.now().strftime('%Y%m%d')}-{hash(application_data.get('법인명(국문)', 'Unknown')) % 10000:04d}"
    doc.add_paragraph(f"견적서 번호: {quotation_number}")
    doc.add_paragraph(f"견적일: {datetime.now().strftime('%Y년 %m월 %d일')}")
    
    # 유효기간
    doc.add_paragraph(f"견적 유효기간: {datetime.now().strftime('%Y년 %m월 %d일')}부터 30일간")
    
    # 바이트 버퍼로 변환
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    return buffer

@app.route('/health', methods=['GET'])
def health_check():
    """헬스 체크 엔드포인트"""
    return jsonify({'status': 'healthy', 'message': 'Quotation API is running'})

@app.route('/generate-quotation', methods=['POST'])
def generate_quotation():
    """견적서 생성 API"""
    try:
        data = request.json
        application_data = data['applicationData']
        
        print(f"견적서 생성 요청: {application_data.get('법인명(국문)', 'Unknown')}")
        
        # 간단한 Word 문서 생성
        docx_buffer = create_simple_quotation_docx(application_data)
        
        # 파일명 생성
        filename = f"LRQA_견적서_{application_data.get('법인명(국문)', 'Unknown')}_{datetime.now().strftime('%Y%m%d')}.docx"
        
        return send_file(
            docx_buffer,
            as_attachment=True,
            download_name=filename,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        
    except Exception as e:
        print(f"오류 발생: {str(e)}")
        import traceback
        traceback.print_exc()
        return jsonify({'error': str(e), 'message': '견적서 생성 중 오류가 발생했습니다.'}), 500

if __name__ == '__main__':
    print("Flask 서버 시작 중...")
    print("헬스 체크: http://localhost:5000/health")
    print("견적서 생성: http://localhost:5000/generate-quotation")
    app.run(debug=False, host='127.0.0.1', port=5000)
